"""
ZettaBrain Lite — FastAPI Web Server
Single-user RAG + Skills platform with multi-provider LLM support.
"""

import asyncio
import json
import os
import re
import shutil
import subprocess
import sys
import threading
import time
from pathlib import Path
from typing import Optional

import requests
from fastapi import FastAPI, HTTPException, Request, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from .config import (
    BASE_DIR,
    CHROMA_DIR,
    DATA_DIR,
    EMBED_MODEL,
    LLM_MODEL,
    OLLAMA_HOST,
    SKILLS_DIR,
    STORAGE_CONF,
    get_setting,
    load_config,
    save_config,
    set_setting,
)
from .retrieval import ADVANCED_RAG_PROMPT, advanced_retrieve, format_context, hybrid_retrieve

PKG_DIR = Path(__file__).parent
STATIC_DIR = PKG_DIR / "static"
CHROMA_PATH = CHROMA_DIR / "default"
INGEST_LOG = DATA_DIR / "ingested_files.json"

# ── Vectorstore cache ────────────────────────────────────────────────────────
_vs_lock = threading.Lock()
_vs_cache: dict = {"vs": None}


def _get_vs():
    with _vs_lock:
        if _vs_cache["vs"] is None:
            from langchain_chroma import Chroma
            from langchain_ollama import OllamaEmbeddings

            cfg = load_config()
            embed_provider = cfg.get("embed_provider", "ollama")
            embed_model = cfg.get("embed_model", EMBED_MODEL)
            ollama_host = cfg.get("ollama_host", OLLAMA_HOST)

            if embed_provider == "ollama":
                emb = OllamaEmbeddings(model=embed_model, base_url=ollama_host)
            elif embed_provider == "openai":
                from langchain_openai import OpenAIEmbeddings

                emb = OpenAIEmbeddings(model=embed_model, api_key=cfg.get("openai_api_key"))
            else:
                emb = OllamaEmbeddings(model=embed_model, base_url=ollama_host)

            _vs_cache["vs"] = Chroma(
                persist_directory=str(CHROMA_PATH),
                embedding_function=emb,
                collection_name="zettabrain_docs",
            )
        return _vs_cache["vs"]


def _reset_vs_cache():
    with _vs_lock:
        _vs_cache["vs"] = None


# ── GPU detection ────────────────────────────────────────────────────────────
def _get_gpu_info() -> dict:
    info = {"type": "none", "name": None, "vram_gb": 0, "recommended_model": None}
    try:
        out = (
            subprocess.check_output(
                ["nvidia-smi", "--query-gpu=name,memory.total", "--format=csv,noheader,nounits"],
                stderr=subprocess.DEVNULL,
                timeout=5,
            )
            .decode()
            .strip()
            .splitlines()[0]
        )
        name, vram_mb = [s.strip() for s in out.split(",")]
        vram_gb = int(vram_mb) // 1024
        info.update({"type": "nvidia", "name": name, "vram_gb": vram_gb})
    except Exception:
        pass

    if info["type"] == "none":
        try:
            out = subprocess.check_output(
                ["rocm-smi", "--showmeminfo", "vram"], stderr=subprocess.DEVNULL, timeout=5
            ).decode()
            import re

            m = re.search(r"(\d+)\s*MB", out)
            if m:
                info.update({"type": "amd", "vram_gb": int(m.group(1)) // 1024})
        except Exception:
            pass

    if info["type"] == "none":
        try:
            if subprocess.check_output(["uname", "-m"]).decode().strip() == "arm64":
                mem = int(subprocess.check_output(["sysctl", "-n", "hw.memsize"]).decode().strip())
                info.update({"type": "apple_silicon", "vram_gb": mem // 1_073_741_824})
        except Exception:
            pass

    vg = info["vram_gb"]
    if info["type"] == "none" or vg == 0:
        info["recommended_model"] = "phi4-mini"
    elif vg >= 24:
        info["recommended_model"] = "qwen2.5:32b"
    elif vg >= 16:
        info["recommended_model"] = "qwen2.5:14b"
    elif vg >= 12:
        info["recommended_model"] = "mistral-nemo:12b"
    elif vg >= 8:
        info["recommended_model"] = "llama3.1:8b"
    elif vg >= 5:
        info["recommended_model"] = "mistral:7b"
    else:
        info["recommended_model"] = "phi4-mini"

    return info


# ── Helpers ──────────────────────────────────────────────────────────────────
def _ollama_running() -> bool:
    ollama_url = get_setting("ollama_host") or OLLAMA_HOST
    _url = ollama_url.replace("https://", "http://")
    try:
        r = requests.get(_url, timeout=5)
        return r.status_code == 200
    except Exception:
        return False


def _get_chunk_count() -> int:
    try:
        import chromadb

        client = chromadb.PersistentClient(path=str(CHROMA_PATH))
        return client.get_collection("zettabrain_docs").count()
    except Exception:
        return 0


def _get_sources() -> list:
    if not INGEST_LOG.exists():
        return []
    try:
        data = json.loads(INGEST_LOG.read_text(encoding="utf-8"))
        return sorted([Path(p).name for p in data.keys()])
    except Exception:
        return []


def _get_ollama_models() -> list:
    ollama_url = get_setting("ollama_host") or OLLAMA_HOST
    try:
        r = requests.get(f"{ollama_url}/api/tags", timeout=5)
        if r.status_code == 200:
            return [m["name"] for m in r.json().get("models", [])]
    except Exception:
        pass
    return []


def _get_available_models() -> list:
    """Return all available models with provider prefix for the model dropdown."""
    models = []

    # Trial model (first in list if available)
    try:
        from .trial import get_trial_model_entry

        trial_entry = get_trial_model_entry()
        if trial_entry:
            models.append(trial_entry)
    except Exception:
        pass

    # Local Ollama models
    for m in _get_ollama_models():
        models.append({"id": f"ollama:{m}", "label": f"Local-{m}", "provider": "ollama"})

    # Cloud models from config
    cfg = load_config()

    if cfg.get("groq_api_key"):
        groq_model = cfg.get("groq_model", "llama-3.1-8b-instant")
        models.append({"id": f"groq:{groq_model}", "label": f"Groq-{groq_model}", "provider": "groq"})

    if cfg.get("openai_api_key"):
        oai_model = cfg.get("openai_model", "gpt-4o-mini")
        models.append({"id": f"openai:{oai_model}", "label": f"OpenAI-{oai_model}", "provider": "openai"})

    if cfg.get("anthropic_api_key"):
        claude_model = cfg.get("claude_model", "claude-sonnet-4-6")
        models.append({"id": f"claude:{claude_model}", "label": f"Claude-{claude_model}", "provider": "claude"})

    if cfg.get("together_api_key"):
        tog_model = cfg.get("together_model", "meta-llama/Meta-Llama-3.1-8B-Instruct-Turbo")
        short = tog_model.split("/")[-1] if "/" in tog_model else tog_model
        models.append({"id": f"together:{tog_model}", "label": f"Together-{short}", "provider": "together"})

    if cfg.get("cerebras_api_key"):
        cer_model = cfg.get("cerebras_model", "llama3.1-8b")
        models.append({"id": f"cerebras:{cer_model}", "label": f"Cerebras-{cer_model}", "provider": "cerebras"})

    if cfg.get("openrouter_api_key"):
        or_model = cfg.get("openrouter_model", "meta-llama/llama-3.1-8b-instruct:free")
        short = or_model.split("/")[-1] if "/" in or_model else or_model
        models.append({"id": f"openrouter:{or_model}", "label": f"OpenRouter-{short}", "provider": "openrouter"})

    if cfg.get("fireworks_api_key"):
        fw_model = cfg.get("fireworks_model", "accounts/fireworks/models/llama-v3p1-8b-instruct")
        short = fw_model.split("/")[-1] if "/" in fw_model else fw_model
        models.append({"id": f"fireworks:{fw_model}", "label": f"Fireworks-{short}", "provider": "fireworks"})

    if cfg.get("gemini_api_key"):
        gem_model = cfg.get("gemini_model", "gemini-3.5-flash-lite")
        models.append({"id": f"gemini:{gem_model}", "label": f"Gemini-{gem_model}", "provider": "gemini"})

    return models


def _get_storage_sources() -> list:
    sources = []
    if STORAGE_CONF.exists():
        for line in STORAGE_CONF.read_text(encoding="utf-8").splitlines():
            if line.startswith("#") or not line.strip():
                continue
            parts = line.split("|")
            if len(parts) >= 4:
                sources.append(
                    {
                        "role": parts[0].strip(),
                        "type": parts[1].strip(),
                        "label": parts[2].strip(),
                        "path": parts[3].strip(),
                    }
                )
    if not sources:
        cfg = load_config()
        docs_folder = cfg.get("docs_folder", str(DATA_DIR / "documents"))
        sources.append({"role": "primary", "type": "local", "label": "Documents", "path": docs_folder})
    return sources


def _count_docs_all_sources() -> int:
    total = 0
    seen = set()
    for src in _get_storage_sources():
        p = Path(src["path"])
        if p in seen or not p.exists():
            continue
        seen.add(p)
        for ext in ["*.pdf", "*.txt", "*.docx", "*.md"]:
            total += len(list(p.rglob(ext)))
    return total


def _resolve_llm_for_chat(model_id: str):
    """Parse a model_id like 'ollama:llama3.1:8b' or 'groq:llama-3.1-8b-instant'
    and return a (provider, model, api_key, ollama_host) tuple."""
    cfg = load_config()
    if ":" in model_id:
        parts = model_id.split(":", 1)
        provider = parts[0]
        model = parts[1]
    else:
        provider = "ollama"
        model = model_id

    ollama_host = cfg.get("ollama_host", OLLAMA_HOST)
    api_key = None

    if provider == "ollama":
        pass
    elif provider == "trial":
        pass
    elif provider == "openai":
        api_key = cfg.get("openai_api_key")
    elif provider == "claude":
        api_key = cfg.get("anthropic_api_key")
    elif provider == "gemini":
        api_key = cfg.get("gemini_api_key") or os.environ.get("GEMINI_API_KEY")
    else:
        api_key = cfg.get(f"{provider}_api_key")

    return provider, model, api_key, ollama_host


# ── App ──────────────────────────────────────────────────────────────────────
app = FastAPI(title="ZettaBrain Lite", version="0.1.0")


@app.on_event("startup")
async def _warmup():
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    CHROMA_PATH.mkdir(parents=True, exist_ok=True)
    SKILLS_DIR.mkdir(parents=True, exist_ok=True)
    if CHROMA_PATH.exists():
        try:
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, _get_vs)
        except Exception:
            pass


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


# ── Request Models ───────────────────────────────────────────────────────────
class ChatRequest(BaseModel):
    question: str
    model: Optional[str] = None


class IngestRequest(BaseModel):
    folder: Optional[str] = None
    rebuild: bool = False


class SettingsUpdate(BaseModel):
    settings: dict


class PullRequest(BaseModel):
    model: str


class StorageAddRequest(BaseModel):
    type: str  # local, nfs, smb, s3
    label: str
    role: str = "secondary"
    # Local
    path: Optional[str] = None
    # NFS
    server_ip: Optional[str] = None
    export_path: Optional[str] = None
    mount_point: Optional[str] = None
    nfs_version: Optional[str] = "4"
    # SMB
    share_name: Optional[str] = None
    username: Optional[str] = None
    password: Optional[str] = None
    domain: Optional[str] = None
    # S3
    bucket: Optional[str] = None
    region: Optional[str] = None
    access_key: Optional[str] = None
    secret_key: Optional[str] = None


class GenerateBody(BaseModel):
    input: str
    skill_name: str
    model: Optional[str] = None
    context: dict = {}
    temperature: Optional[float] = None
    max_tokens: Optional[int] = None


class SkillUploadBody(BaseModel):
    content: str
    filename: str


class SkillDraftBody(BaseModel):
    goal: str
    name: str = ""
    sections: list = []
    tone: list = ["Professional"]
    requires_corpus: bool = False
    citations: bool = False
    max_tokens: int = 2000
    example_output: str = ""
    model: Optional[str] = None


# ── Routes: UI ───────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def root():
    index = STATIC_DIR / "index.html"
    if index.exists():
        return HTMLResponse(index.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>ZettaBrain Lite</h1><p>Static files not found.</p>")


# ── Routes: Status & Models ──────────────────────────────────────────────────
@app.get("/api/status")
async def status():
    return {
        "ollama": {
            "running": _ollama_running(),
            "url": get_setting("ollama_host") or OLLAMA_HOST,
            "models": _get_ollama_models(),
            "active_llm": get_setting("llm_model") or LLM_MODEL,
            "active_embed": get_setting("embed_model") or EMBED_MODEL,
        },
        "vectorstore": {
            "exists": CHROMA_PATH.exists(),
            "chunks": _get_chunk_count(),
            "path": str(CHROMA_PATH),
        },
        "storage": {
            "doc_count": _count_docs_all_sources(),
            "sources": _get_storage_sources(),
        },
        "hardware": _get_gpu_info(),
        "sources": _get_sources(),
    }


@app.get("/api/models")
async def get_models():
    return {"models": _get_available_models()}


@app.get("/api/trial")
async def trial_status():
    try:
        from .trial import get_trial_usage

        return get_trial_usage()
    except Exception:
        return {"requests_used": 0, "requests_remaining": 0, "max_requests": 5, "exhausted": True}


@app.get("/api/sources")
async def get_sources():
    return {"sources": _get_sources()}


# ── Routes: Settings ─────────────────────────────────────────────────────────
@app.get("/api/settings")
async def get_settings():
    cfg = load_config()
    safe_cfg = {}
    for k, v in cfg.items():
        if "key" in k.lower() or "secret" in k.lower() or "password" in k.lower():
            safe_cfg[k] = "***" + str(v)[-4:] if v else ""
        else:
            safe_cfg[k] = v
    return {"settings": safe_cfg}


@app.post("/api/settings")
async def update_settings(body: SettingsUpdate):
    cfg = load_config()
    for k, v in body.settings.items():
        if v and not (isinstance(v, str) and v.startswith("***")):
            cfg[k] = v
    save_config(cfg)
    _reset_vs_cache()
    return {"success": True}


@app.post("/api/settings/logo")
async def upload_logo(request: Request):
    """Upload organization logo (accepts multipart form or base64 JSON)."""
    content_type = request.headers.get("content-type", "")

    if "multipart" in content_type:
        form = await request.form()
        file = form.get("logo")
        if not file:
            raise HTTPException(status_code=400, detail="No logo file provided")
        logo_bytes = await file.read()
        filename = file.filename or "logo.png"
    else:
        import base64

        body = await request.json()
        logo_bytes = base64.b64decode(body.get("data", ""))
        filename = body.get("filename", "logo.png")

    logo_dir = DATA_DIR / "branding"
    logo_dir.mkdir(parents=True, exist_ok=True)

    ext = Path(filename).suffix or ".png"
    logo_path = logo_dir / f"logo{ext}"
    logo_path.write_bytes(logo_bytes)

    set_setting("logo_path", str(logo_path))
    return {"success": True, "path": str(logo_path)}


@app.get("/api/settings/logo")
async def get_logo():
    """Serve the uploaded logo."""
    logo_path_str = get_setting("logo_path")
    if not logo_path_str:
        raise HTTPException(status_code=404, detail="No logo uploaded")

    logo_path = Path(logo_path_str)
    if not logo_path.exists():
        raise HTTPException(status_code=404, detail="Logo file not found")

    from fastapi.responses import FileResponse

    return FileResponse(logo_path)


# ── Routes: Storage Management ───────────────────────────────────────────────
@app.get("/api/storage")
async def list_storage():
    return {"sources": _get_storage_sources(), "doc_count": _count_docs_all_sources()}


def _validate_ip(ip: str) -> bool:
    return bool(re.match(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$", ip))


def _is_macos() -> bool:
    return sys.platform == "darwin"


def _mount_nfs(server_ip: str, export_path: str, mount_point: str, nfs_version: str = "4") -> str:
    Path(mount_point).mkdir(parents=True, exist_ok=True)
    if _is_macos():
        opts = f"nfsvers={nfs_version},rsize=1048576,wsize=1048576,hard,timeo=600,retrans=2"
        cmd = ["mount", "-t", "nfs", "-o", opts, f"{server_ip}:{export_path}", mount_point]
    else:
        opts = f"defaults,_netdev,nfsvers={nfs_version},rsize=1048576,wsize=1048576,hard,timeo=600,retrans=2"
        cmd = ["mount", "-t", "nfs", "-o", opts, f"{server_ip}:{export_path}", mount_point]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
    if result.returncode != 0:
        return result.stderr.strip() or f"Mount failed (exit code {result.returncode})"
    return ""


def _mount_smb(
    server_ip: str, share_name: str, mount_point: str, username: str = "guest", password: str = "", domain: str = ""
) -> str:
    Path(mount_point).mkdir(parents=True, exist_ok=True)
    creds_dir = Path("/etc/zettabrain")
    creds_dir.mkdir(parents=True, exist_ok=True)
    creds_file = creds_dir / f"smb-{server_ip}.credentials"
    creds_content = f"username={username}\npassword={password}\n"
    if domain:
        creds_content += f"domain={domain}\n"
    creds_file.write_text(creds_content, encoding="utf-8")
    creds_file.chmod(0o600)

    if _is_macos():
        mount_url = f"//{username}:{password}@{server_ip}/{share_name}"
        cmd = ["mount", "-t", "smbfs", mount_url, mount_point]
    else:
        opts = f"uid=0,gid=0,file_mode=0755,dir_mode=0755,noperm,_netdev,credentials={creds_file}"
        cmd = ["mount", "-t", "cifs", f"//{server_ip}/{share_name}", mount_point, "-o", opts]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
    if result.returncode != 0:
        return result.stderr.strip() or f"Mount failed (exit code {result.returncode})"
    return ""


@app.post("/api/storage")
async def add_storage(body: StorageAddRequest):
    error = ""
    final_path = ""

    if body.type == "local":
        if not body.path:
            raise HTTPException(status_code=400, detail="Path is required for local storage")
        p = Path(body.path)
        p.mkdir(parents=True, exist_ok=True)
        final_path = body.path

    elif body.type == "nfs":
        if not body.server_ip or not body.export_path or not body.mount_point:
            raise HTTPException(status_code=400, detail="NFS requires server_ip, export_path, and mount_point")
        if not _validate_ip(body.server_ip):
            raise HTTPException(status_code=400, detail="Invalid IP address format")
        if not body.export_path.startswith("/"):
            raise HTTPException(status_code=400, detail="Export path must start with /")
        error = _mount_nfs(body.server_ip, body.export_path, body.mount_point, body.nfs_version or "4")
        final_path = body.mount_point

    elif body.type == "smb":
        if not body.server_ip or not body.share_name or not body.mount_point:
            raise HTTPException(status_code=400, detail="SMB requires server_ip, share_name, and mount_point")
        if not _validate_ip(body.server_ip):
            raise HTTPException(status_code=400, detail="Invalid IP address format")
        error = _mount_smb(
            body.server_ip,
            body.share_name,
            body.mount_point,
            body.username or "guest",
            body.password or "",
            body.domain or "",
        )
        final_path = body.mount_point

    elif body.type == "s3":
        if not body.bucket or not body.mount_point:
            raise HTTPException(status_code=400, detail="S3 requires bucket and mount_point")
        Path(body.mount_point).mkdir(parents=True, exist_ok=True)
        s3_cfg = BASE_DIR / "s3_credentials"
        s3_cfg.parent.mkdir(parents=True, exist_ok=True)
        s3_content = f"{body.access_key or ''}:{body.secret_key or ''}"
        s3_cfg.write_text(s3_content, encoding="utf-8")
        s3_cfg.chmod(0o600)
        final_path = body.mount_point
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported storage type: {body.type}")

    if error:
        raise HTTPException(status_code=500, detail=f"Mount failed: {error}")

    STORAGE_CONF.parent.mkdir(parents=True, exist_ok=True)
    line = f"{body.role}|{body.type}|{body.label}|{final_path}\n"
    with open(STORAGE_CONF, "a", encoding="utf-8") as f:
        f.write(line)

    return {"success": True, "source": {"role": body.role, "type": body.type, "label": body.label, "path": final_path}}


@app.post("/api/storage/test")
async def test_storage(body: StorageAddRequest):
    """Test storage connectivity without mounting."""
    if body.type == "local":
        exists = Path(body.path or "").exists() if body.path else False
        return {"reachable": exists, "message": "Path exists" if exists else "Path not found"}

    if body.type == "nfs":
        if not body.server_ip:
            return {"reachable": False, "message": "Server IP required"}
        result = subprocess.run(
            ["ping", "-c", "1", "-W", "3", body.server_ip],
            capture_output=True,
            text=True,
            timeout=10,
        )
        reachable = result.returncode == 0
        return {"reachable": reachable, "message": f"Server {'reachable' if reachable else 'unreachable'}"}

    if body.type == "smb":
        if not body.server_ip:
            return {"reachable": False, "message": "Server IP required"}
        result = subprocess.run(
            ["ping", "-c", "1", "-W", "3", body.server_ip],
            capture_output=True,
            text=True,
            timeout=10,
        )
        reachable = result.returncode == 0
        return {"reachable": reachable, "message": f"Server {'reachable' if reachable else 'unreachable'}"}

    return {"reachable": False, "message": "Test not supported for this type"}


@app.delete("/api/storage/{index}")
async def remove_storage(index: int):
    if not STORAGE_CONF.exists():
        raise HTTPException(status_code=404, detail="No storage sources configured")
    lines = [
        line
        for line in STORAGE_CONF.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.startswith("#")
    ]
    if index < 0 or index >= len(lines):
        raise HTTPException(status_code=404, detail="Source not found")
    lines.pop(index)
    STORAGE_CONF.write_text("\n".join(lines) + "\n" if lines else "", encoding="utf-8")
    return {"success": True}


# ── Routes: Model Pull ───────────────────────────────────────────────────────
@app.post("/api/pull")
async def pull_model(req: PullRequest):
    ollama_url = get_setting("ollama_host") or OLLAMA_HOST
    queue: asyncio.Queue = asyncio.Queue()
    loop = asyncio.get_event_loop()

    def _do_pull():
        try:
            resp = requests.post(
                f"{ollama_url}/api/pull",
                json={"name": req.model, "stream": True},
                stream=True,
                timeout=600,
            )
            if resp.status_code != 200:
                asyncio.run_coroutine_threadsafe(queue.put(f"Error: Ollama returned {resp.status_code}\n"), loop)
                return
            for line in resp.iter_lines():
                if not line:
                    continue
                try:
                    chunk = json.loads(line)
                    status_msg = chunk.get("status", "")
                    if chunk.get("total") and chunk.get("completed") is not None:
                        pct = int(chunk["completed"] / chunk["total"] * 100)
                        msg = f"{status_msg} {pct}%\n"
                    elif status_msg:
                        msg = f"{status_msg}\n"
                    else:
                        continue
                    asyncio.run_coroutine_threadsafe(queue.put(msg), loop)
                except Exception:
                    pass
        except Exception as exc:
            asyncio.run_coroutine_threadsafe(queue.put(f"Error: {exc}\n"), loop)
        finally:
            asyncio.run_coroutine_threadsafe(queue.put(None), loop)

    loop.run_in_executor(None, _do_pull)

    async def _gen():
        while True:
            msg = await queue.get()
            if msg is None:
                break
            yield msg

    return StreamingResponse(_gen(), media_type="text/plain")


# ── Routes: Ingestion ────────────────────────────────────────────────────────
@app.post("/api/ingest")
async def ingest(req: IngestRequest):
    script = PKG_DIR / "scripts" / "05_ingest_documents.py"
    if not script.exists():
        raise HTTPException(status_code=404, detail="Ingest script not found.")

    DATA_DIR.mkdir(parents=True, exist_ok=True)
    CHROMA_PATH.mkdir(parents=True, exist_ok=True)

    cfg = load_config()
    docs_folder = req.folder or cfg.get("docs_folder")
    if not docs_folder:
        sources = _get_storage_sources()
        docs_folder = sources[0]["path"] if sources else str(DATA_DIR / "documents")

    cmd = [sys.executable, str(script)]
    if req.folder:
        cmd += ["--folder", req.folder]
    if req.rebuild:
        cmd += ["--rebuild"]

    env = os.environ.copy()
    env["ZETTABRAIN_CHROMA"] = str(CHROMA_PATH)
    env["ZETTABRAIN_DOCS"] = docs_folder
    env["OLLAMA_HOST"] = cfg.get("ollama_host", OLLAMA_HOST)

    try:
        result = subprocess.run(
            cmd,
            cwd=str(DATA_DIR),
            capture_output=True,
            text=True,
            timeout=600,
            env=env,
        )
        if result.returncode == 0:
            _reset_vs_cache()
        return {
            "success": result.returncode == 0,
            "output": result.stdout,
            "error": result.stderr if result.returncode != 0 else None,
            "chunks": _get_chunk_count(),
        }
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=408, detail="Ingestion timed out.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/ingest/source/{index}")
async def ingest_source(index: int):
    """Ingest documents from a specific storage source by index."""
    sources = _get_storage_sources()
    if index < 0 or index >= len(sources):
        raise HTTPException(status_code=404, detail="Source not found")
    source = sources[index]
    req = IngestRequest(folder=source["path"])
    return await ingest(req)


# ── Routes: Chat (RAG) ──────────────────────────────────────────────────────
@app.post("/api/chat")
async def chat(req: ChatRequest):
    if _get_chunk_count() == 0:
        raise HTTPException(status_code=422, detail="Vector store is empty. Ingest documents first.")

    model_id = req.model or f"ollama:{get_setting('llm_model') or LLM_MODEL}"
    provider, model, api_key, ollama_host = _resolve_llm_for_chat(model_id)

    question = req.question

    def _run():
        vs = _get_vs()

        # Build a lightweight LLM callable for retrieval stages
        llm_fn = None
        try:
            if provider == "trial":
                from .trial import trial_generate

                def llm_fn(p):
                    return trial_generate(p, temperature=0.0, max_tokens=256)
            elif provider == "ollama":

                def _ollama_fn(p):
                    resp = requests.post(
                        f"{ollama_host}/api/generate",
                        json={"model": model, "prompt": p, "stream": False},
                        timeout=30,
                    )
                    if resp.status_code == 200:
                        return resp.json().get("response", "")
                    return ""

                llm_fn = _ollama_fn
            else:
                from .llm.factory import get_chat_llm as _gcl

                _rag_llm = _gcl(provider=provider, model=model, ollama_host=ollama_host, api_key=api_key)

                def _cloud_fn(p):
                    r = _rag_llm.invoke(p)
                    return r.content.strip() if hasattr(r, "content") else str(r).strip()

                llm_fn = _cloud_fn
        except Exception:
            pass

        t0 = time.monotonic()
        sources = advanced_retrieve(question, vs, llm_fn=llm_fn)
        t_retr = time.monotonic() - t0

        context = format_context(sources)
        prompt_text = ADVANCED_RAG_PROMPT.format(context=context, question=question)

        t1 = time.monotonic()

        if provider == "trial":
            from .trial import trial_generate

            answer = trial_generate(prompt_text, temperature=0.0, max_tokens=1024)
        else:
            from .llm.factory import get_chat_llm

            llm = get_chat_llm(
                provider=provider,
                model=model,
                ollama_host=ollama_host,
                api_key=api_key,
            )

            from langchain_core.prompts import PromptTemplate

            prompt = PromptTemplate.from_template(ADVANCED_RAG_PROMPT)
            response = llm.invoke(prompt.format(context=context, question=question))
            if hasattr(response, "content"):
                answer = response.content.strip()
            else:
                answer = str(response).strip()

        t_gen = time.monotonic() - t1
        return sources, answer, t_retr, t_gen

    try:
        loop = asyncio.get_event_loop()
        sources, answer, t_retr, t_gen = await loop.run_in_executor(None, _run)

        # Save to history
        from .database import ChatHistory, get_session

        with get_session() as session:
            session.add(
                ChatHistory(
                    question=req.question,
                    answer=answer,
                    model=model_id,
                    chunks_searched=len(sources),
                    duration_ms=round((t_retr + t_gen) * 1000),
                    sources=json.dumps([Path(s.metadata.get("source", "?")).name for s in sources]),
                )
            )
            session.commit()

        return {
            "answer": answer,
            "model": model_id,
            "chunks_searched": len(sources),
            "timing": {
                "retrieve_ms": round(t_retr * 1000),
                "generate_ms": round(t_gen * 1000),
            },
            "sources": [
                {
                    "filename": Path(s.metadata.get("source", "?")).name,
                    "page": s.metadata.get("page", ""),
                    "preview": s.page_content[:200],
                }
                for s in sources
            ],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.websocket("/ws/chat")
async def websocket_chat(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            data = await websocket.receive_text()
            payload = json.loads(data)
            question = payload.get("question", "").strip()
            model_id = payload.get("model", f"ollama:{get_setting('llm_model') or LLM_MODEL}")

            if not question:
                await websocket.send_json({"type": "error", "message": "Empty question"})
                continue
            if _get_chunk_count() == 0:
                await websocket.send_json(
                    {"type": "error", "message": "Vector store is empty. Ingest documents first."}
                )
                continue

            provider, model, api_key, ollama_host = _resolve_llm_for_chat(model_id)

            try:
                loop = asyncio.get_event_loop()
                vectorstore = await loop.run_in_executor(None, _get_vs)

                # Build lightweight LLM callable for retrieval stages
                ws_llm_fn = None
                try:
                    if provider == "trial":
                        from .trial import trial_generate as _tg

                        def ws_llm_fn(p):
                            return _tg(p, temperature=0.0, max_tokens=256)
                    elif provider == "ollama":
                        _oh, _m = ollama_host, model

                        def _ws_ollama_fn(p):
                            resp = requests.post(
                                f"{_oh}/api/generate",
                                json={"model": _m, "prompt": p, "stream": False},
                                timeout=30,
                            )
                            return resp.json().get("response", "") if resp.status_code == 200 else ""

                        ws_llm_fn = _ws_ollama_fn
                    else:
                        from .llm.factory import get_chat_llm as _gcl2

                        _ws_llm = _gcl2(provider=provider, model=model, ollama_host=ollama_host, api_key=api_key)

                        def _ws_cloud_fn(p):
                            r = _ws_llm.invoke(p)
                            return r.content.strip() if hasattr(r, "content") else str(r).strip()

                        ws_llm_fn = _ws_cloud_fn
                except Exception:
                    pass

                t_r0 = time.monotonic()
                sources = await loop.run_in_executor(
                    None, lambda: advanced_retrieve(question, vectorstore, llm_fn=ws_llm_fn)
                )
                t_retr = time.monotonic() - t_r0
                source_list = [
                    {
                        "filename": Path(s.metadata.get("source", "?")).name,
                        "page": s.metadata.get("page", ""),
                        "preview": s.page_content[:200],
                    }
                    for s in sources
                ]
                await websocket.send_json({"type": "sources", "sources": source_list})

                context = format_context(sources)
                prompt_text = ADVANCED_RAG_PROMPT.format(context=context, question=question)

                # Stream based on provider
                if provider == "trial":
                    from .trial import trial_generate

                    t_g0 = time.monotonic()
                    full_answer = ""
                    error_msg = None
                    try:
                        result = await loop.run_in_executor(
                            None, lambda: trial_generate(prompt_text, temperature=0.0, max_tokens=1024)
                        )
                        full_answer = result
                        await websocket.send_json({"type": "token", "token": result})
                    except Exception as exc:
                        error_msg = str(exc)

                elif provider == "ollama":
                    queue: asyncio.Queue = asyncio.Queue()
                    t_g0 = time.monotonic()

                    def _stream_ollama():
                        try:
                            resp = requests.post(
                                f"{ollama_host}/api/generate",
                                json={"model": model, "prompt": prompt_text, "stream": True},
                                stream=True,
                                timeout=600,
                            )
                            if resp.status_code != 200:
                                asyncio.run_coroutine_threadsafe(
                                    queue.put(("error", f"Ollama {resp.status_code}")), loop
                                )
                                return
                            for line in resp.iter_lines():
                                if line:
                                    chunk = json.loads(line)
                                    if "error" in chunk:
                                        asyncio.run_coroutine_threadsafe(queue.put(("error", chunk["error"])), loop)
                                        return
                                    token = chunk.get("response", "")
                                    asyncio.run_coroutine_threadsafe(queue.put(("token", token)), loop)
                                    if chunk.get("done"):
                                        break
                        except Exception as exc:
                            asyncio.run_coroutine_threadsafe(queue.put(("error", str(exc))), loop)
                        finally:
                            asyncio.run_coroutine_threadsafe(queue.put(("done", None)), loop)

                    loop.run_in_executor(None, _stream_ollama)

                    full_answer = ""
                    error_msg = None
                    while True:
                        kind, value = await queue.get()
                        if kind == "token":
                            full_answer += value
                            await websocket.send_json({"type": "token", "token": value})
                        elif kind == "error":
                            error_msg = value
                            break
                        else:
                            break

                else:
                    # Cloud providers: use LLM factory streaming
                    from .llm.factory import create_generation_provider

                    t_g0 = time.monotonic()
                    gen_provider = create_generation_provider(
                        provider_name=provider,
                        model=model,
                        api_key=api_key,
                    )
                    full_answer = ""
                    error_msg = None

                    def _stream_cloud():
                        tokens = []
                        for token in gen_provider.stream(prompt_text, temperature=0.0, max_tokens=1024):
                            tokens.append(token)
                        return tokens

                    tokens = await loop.run_in_executor(None, _stream_cloud)
                    for token in tokens:
                        full_answer += token
                        await websocket.send_json({"type": "token", "token": token})

                if error_msg:
                    await websocket.send_json({"type": "error", "message": error_msg})
                    continue

                t_gen = time.monotonic() - t_g0

                # Save to history
                from .database import ChatHistory, get_session

                with get_session() as session:
                    session.add(
                        ChatHistory(
                            question=question,
                            answer=full_answer,
                            model=model_id,
                            chunks_searched=len(sources),
                            duration_ms=round((t_retr + t_gen) * 1000),
                            sources=json.dumps([s["filename"] for s in source_list]),
                        )
                    )
                    session.commit()

                await websocket.send_json(
                    {
                        "type": "done",
                        "answer": full_answer,
                        "model": model_id,
                        "chunks_searched": len(sources),
                        "timing": {
                            "retrieve_ms": round(t_retr * 1000),
                            "generate_ms": round(t_gen * 1000),
                        },
                    }
                )

            except Exception as e:
                msg = str(e)
                if "hnsw" in msg.lower() or "segment" in msg.lower():
                    _reset_vs_cache()
                    msg = "Vector store index is corrupt. Clear it and re-run ingestion."
                await websocket.send_json({"type": "error", "message": msg})

    except WebSocketDisconnect:
        pass


# ── Routes: Skills / Generation ──────────────────────────────────────────────
_BUILTIN_SKILLS_DIR = PKG_DIR / "skills"


@app.get("/api/skills")
async def list_skills():
    from .generation.skill_parser import SkillParser

    seen_names = set()
    skills = []

    for skills_dir in [SKILLS_DIR, _BUILTIN_SKILLS_DIR]:
        if not skills_dir.exists():
            continue
        source = "user" if skills_dir == SKILLS_DIR else "builtin"
        for f in sorted(skills_dir.glob("*.md")):
            try:
                skill = SkillParser.parse_file(f)
                if skill.name not in seen_names:
                    seen_names.add(skill.name)
                    skills.append(
                        {
                            "name": skill.name,
                            "version": skill.version,
                            "description": skill.description,
                            "business_type": skill.business_type,
                            "requires_corpus": skill.requires_corpus,
                            "tags": skill.tags,
                            "source": source,
                        }
                    )
            except Exception:
                continue

    return {"skills": skills}


@app.post("/api/generate")
async def generate_document(body: GenerateBody):
    from .generation.engine import GenerationEngine
    from .generation.models import GenerationRequest
    from .generation.skill_parser import load_skill

    skill_file = _find_skill_file(body.skill_name)
    if not skill_file:
        raise HTTPException(status_code=404, detail=f"Skill '{body.skill_name}' not found")

    skill = load_skill(skill_file)

    corpus_retriever = None
    if _get_chunk_count() > 0:
        corpus_retriever = _build_corpus_retriever()

    model_id = body.model or f"ollama:{get_setting('llm_model') or LLM_MODEL}"
    provider, model, api_key, ollama_host = _resolve_llm_for_chat(model_id)

    if provider == "trial":
        from .llm.base import LLMProvider as _LP
        from .trial import _get_trial_model
        from .trial import trial_generate as _trial_gen

        class _TrialProvider(_LP):
            def generate(self, prompt, temperature=0.7, max_tokens=2000, **kw):
                return _trial_gen(prompt, temperature=temperature, max_tokens=max_tokens)

            def stream(self, prompt, temperature=0.7, max_tokens=2000, **kw):
                yield self.generate(prompt, temperature, max_tokens)

            def check_health(self):
                return True

            def get_model_info(self):
                return {"provider": "trial", "model": _get_trial_model()}

        llm_provider = _TrialProvider()
    else:
        from .llm.factory import create_generation_provider

        gen_kwargs = {"provider_name": provider, "model": model}
        if api_key:
            gen_kwargs["api_key"] = api_key
        if provider == "ollama":
            gen_kwargs["base_url"] = ollama_host
        llm_provider = create_generation_provider(**gen_kwargs)

    engine = GenerationEngine(llm_provider=llm_provider, corpus_retriever=corpus_retriever)

    request = GenerationRequest(
        input=body.input,
        skill_name=body.skill_name,
        context=body.context,
        temperature=body.temperature,
        max_tokens=body.max_tokens,
    )

    result = engine.generate(skill, request)

    if not result.success:
        raise HTTPException(status_code=500, detail=f"Generation failed: {result.error}")

    # Save to history
    from .database import GenerationHistory, get_session

    with get_session() as session:
        record = GenerationHistory(
            skill_name=result.skill_name,
            skill_version=result.skill_version,
            input_text=body.input,
            output_content=result.content,
            citations=json.dumps(result.citations) if result.citations else None,
            generation_time_ms=result.generation_time_ms,
            metadata_json=json.dumps(result.metadata) if result.metadata else None,
        )
        session.add(record)
        session.commit()
        session.refresh(record)

    return {
        "id": record.id,
        "content": result.content,
        "skill_name": result.skill_name,
        "skill_version": result.skill_version,
        "citations": result.citations,
        "generation_time_ms": result.generation_time_ms,
    }


@app.post("/api/skills/draft")
async def draft_skill(body: SkillDraftBody):
    import asyncio

    from .skill_drafter import extract_rules, generate_skill_draft

    model_id = body.model or f"ollama:{get_setting('llm_model') or LLM_MODEL}"
    provider, model, api_key, ollama_host = _resolve_llm_for_chat(model_id)

    if provider == "trial":
        from .trial import trial_generate as _trial_gen

        def llm_fn(prompt: str) -> str:
            return _trial_gen(prompt, temperature=0.4, max_tokens=4000)
    elif provider == "ollama":
        import requests as _req

        def llm_fn(prompt: str) -> str:
            r = _req.post(
                f"{ollama_host}/api/generate",
                json={"model": model, "prompt": prompt, "stream": False,
                      "options": {"temperature": 0.4, "num_predict": 4000}},
                timeout=120,
            )
            r.raise_for_status()
            return r.json().get("response", "")
    else:
        from .llm.factory import create_generation_provider

        gen_kwargs = {"provider_name": provider, "model": model}
        if api_key:
            gen_kwargs["api_key"] = api_key
        _provider = create_generation_provider(**gen_kwargs)

        def llm_fn(prompt: str) -> str:
            return _provider.generate(prompt, temperature=0.4, max_tokens=4000)

    rules: list = []
    doc_types: list = []
    if body.requires_corpus and _get_chunk_count() > 0:
        try:
            retriever = _build_corpus_retriever()
            rules = await asyncio.to_thread(extract_rules, llm_fn, retriever)
        except Exception:
            rules = []
        sources = _get_sources()
        doc_types = list({Path(s).suffix.lstrip(".").lower() for s in sources if "." in s})

    try:
        result = await asyncio.to_thread(
            generate_skill_draft,
            llm_fn=llm_fn,
            goal=body.goal,
            name=body.name,
            sections=body.sections,
            tone=body.tone,
            requires_corpus=body.requires_corpus,
            citations=body.citations,
            max_tokens=body.max_tokens,
            example_output=body.example_output,
            rules=rules,
        )
    except ValueError as e:
        raise HTTPException(status_code=502, detail=str(e))
    except Exception:
        raise HTTPException(
            status_code=502,
            detail="Skill drafting failed. The model may be slow or unreachable — try again or choose a different model.",
        )

    result["doc_types"] = doc_types
    return result


@app.post("/api/skills/upload")
async def upload_skill(body: SkillUploadBody):
    import tempfile

    from .generation.skill_parser import SkillParser

    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as tmp:
        tmp.write(body.content)
        tmp_path = tmp.name

    try:
        skill = SkillParser.parse_and_validate(tmp_path)
    except (ValueError, FileNotFoundError) as e:
        Path(tmp_path).unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    SKILLS_DIR.mkdir(parents=True, exist_ok=True)
    filename = body.filename if body.filename.endswith(".md") else f"{body.filename}.md"
    skill_path = SKILLS_DIR / filename
    skill_path.write_text(body.content, encoding="utf-8")

    from .skill_drafter import validate_skill

    quality = validate_skill(body.content).as_dict()

    return {
        "message": f"Skill '{skill.name}' uploaded",
        "skill": {
            "name": skill.name,
            "version": skill.version,
            "description": skill.description,
        },
        "quality": quality,
    }


@app.delete("/api/skills/{skill_name}")
async def delete_skill(skill_name: str):
    from .generation.skill_parser import SkillParser

    for skills_dir in [SKILLS_DIR, _BUILTIN_SKILLS_DIR]:
        if not skills_dir.exists():
            continue
        for f in skills_dir.glob("*.md"):
            try:
                skill = SkillParser.parse_file(f)
                if skill.name == skill_name:
                    f.unlink()
                    return {"message": f"Skill '{skill_name}' deleted"}
            except Exception:
                continue
    raise HTTPException(status_code=404, detail="Skill not found")


@app.get("/api/skills/{skill_name}/content")
async def get_skill_content(skill_name: str):
    skill_file = _find_skill_file(skill_name)
    if not skill_file:
        raise HTTPException(status_code=404, detail="Skill not found")
    content = Path(skill_file).read_text(encoding="utf-8")
    return {"name": skill_name, "content": content, "filename": Path(skill_file).name}


# ── Routes: History ──────────────────────────────────────────────────────────
@app.get("/api/history/chat")
async def chat_history(limit: int = 50):
    from sqlmodel import select

    from .database import ChatHistory, get_session

    with get_session() as session:
        rows = session.exec(select(ChatHistory).order_by(ChatHistory.created_at.desc()).limit(limit)).all()
    return [
        {
            "id": r.id,
            "question": r.question[:200],
            "answer": r.answer[:300],
            "model": r.model,
            "created_at": r.created_at.isoformat(),
        }
        for r in rows
    ]


@app.get("/api/history/generation")
async def generation_history(limit: int = 50):
    from sqlmodel import select

    from .database import GenerationHistory, get_session

    with get_session() as session:
        rows = session.exec(select(GenerationHistory).order_by(GenerationHistory.created_at.desc()).limit(limit)).all()
    return [
        {
            "id": r.id,
            "skill_name": r.skill_name,
            "input_text": r.input_text[:200],
            "output_preview": r.output_content[:300],
            "created_at": r.created_at.isoformat(),
        }
        for r in rows
    ]


@app.get("/api/history/generation/{record_id}")
async def generation_detail(record_id: int):
    from .database import GenerationHistory, get_session

    with get_session() as session:
        record = session.get(GenerationHistory, record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")
        return {
            "id": record.id,
            "skill_name": record.skill_name,
            "skill_version": record.skill_version,
            "input_text": record.input_text,
            "output_content": record.output_content,
            "citations": json.loads(record.citations) if record.citations else [],
            "generation_time_ms": record.generation_time_ms,
            "metadata": json.loads(record.metadata_json) if record.metadata_json else {},
            "created_at": record.created_at.isoformat(),
        }


def _parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Parse a markdown table starting at `start`. Returns (rows, next_line_index)."""
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped.startswith("|"):
            break
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if not all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def _render_pdf_table(pdf, rows: list[list[str]], lm: float, pw: float) -> None:
    """Render a parsed markdown table into the PDF with borders and shaded header."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    for row in rows:
        while len(row) < n_cols:
            row.append("")

    col_widths = []
    for c in range(n_cols):
        max_len = max(len(row[c]) for row in rows)
        col_widths.append(max(max_len, 4))
    total = sum(col_widths)
    col_widths = [w / total * pw for w in col_widths]

    row_h = 7
    pdf.set_draw_color(200, 200, 200)
    pdf.set_line_width(0.3)

    for r_idx, row in enumerate(rows):
        if pdf.get_y() + row_h > pdf.h - pdf.b_margin:
            pdf.add_page()

        x = lm
        if r_idx == 0:
            pdf.set_fill_color(240, 245, 250)
            pdf.set_font("Helvetica", "B", 9)
        else:
            pdf.set_fill_color(255, 255, 255)
            pdf.set_font("Helvetica", "", 9)

        is_total_row = any("total" in cell.lower() for cell in row) and r_idx == len(rows) - 1
        if is_total_row:
            pdf.set_fill_color(245, 248, 255)
            pdf.set_font("Helvetica", "B", 9)

        for c_idx, cell in enumerate(row):
            pdf.set_xy(x, pdf.get_y())
            pdf.cell(col_widths[c_idx], row_h, cell[:50], border=1, fill=True)
            x += col_widths[c_idx]
        pdf.ln(row_h)

    pdf.ln(3)


@app.get("/api/export/{record_id}/pdf")
async def export_pdf(record_id: int):
    from .database import GenerationHistory, get_session

    with get_session() as session:
        record = session.get(GenerationHistory, record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="fpdf2 not installed. Install with: pip install fpdf2")

    org_name = get_setting("org_name") or "Organization"
    logo_path_str = get_setting("logo_path")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    if logo_path_str and Path(logo_path_str).exists():
        try:
            pdf.image(logo_path_str, x=15, y=10, h=18)
        except Exception:
            pass

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_xy(15, 32)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 10, org_name, new_x="LMARGIN", new_y="NEXT")

    pdf.set_draw_color(59, 130, 246)
    pdf.set_line_width(0.5)
    pdf.line(15, pdf.get_y() + 2, 195, pdf.get_y() + 2)
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 116, 139)
    pdf.cell(
        0,
        6,
        f"Skill: {record.skill_name}  |  Generated: {record.created_at.strftime('%B %d, %Y')}",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(6)

    pdf.set_text_color(30, 41, 59)
    content = record.output_content or ""

    def _sanitize_for_pdf(text):
        replacements = {
            "•": "-",
            "–": "-",
            "—": "-",
            "‘": "'",
            "’": "'",
            "“": '"',
            "”": '"',
            "…": "...",
            " ": " ",
            "‒": "-",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    content = _sanitize_for_pdf(content)

    lm = pdf.l_margin
    pw = pdf.w - pdf.l_margin - pdf.r_margin

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            pdf.ln(3)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = _parse_md_table(lines, i)
            _render_pdf_table(pdf, rows, lm, pw)
            continue

        pdf.set_x(lm)

        if stripped.startswith("# "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_x(lm)
            pdf.multi_cell(pw, 8, stripped[2:])
            pdf.ln(2)
        elif stripped.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_x(lm)
            pdf.multi_cell(pw, 7, stripped[3:])
            pdf.ln(2)
        elif stripped.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_x(lm)
            pdf.multi_cell(pw, 6, stripped[4:])
            pdf.ln(1)
        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
            pdf.set_font("Helvetica", "", 10)
            indent = 8
            pdf.set_x(lm + indent)
            pdf.multi_cell(pw - indent, 5.5, "- " + stripped[2:].strip())
        elif stripped.startswith("**") and stripped.endswith("**"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(pw, 5.5, stripped.strip("*"))
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(pw, 5.5, stripped.replace("**", ""))
        i += 1

    pdf.ln(10)
    pdf.set_draw_color(200, 200, 200)
    pdf.set_line_width(0.3)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, f"Generated by ZettaBrain  |  {record.created_at.strftime('%B %d, %Y at %I:%M %p')}", align="C")

    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        pdf.output(tmp.name)
        tmp_path = tmp.name

    from fastapi.responses import FileResponse

    safe_name = record.skill_name.replace(" ", "_")
    return FileResponse(
        tmp_path,
        media_type="application/pdf",
        filename=f"{safe_name}_{record.created_at.strftime('%Y%m%d')}.pdf",
    )


def _render_docx_table(doc, rows: list[list[str]]) -> None:
    """Render a parsed markdown table into the DOCX with borders and shaded header."""
    from docx.oxml.ns import nsdecls
    from docx.oxml.parser import parse_xml
    from docx.shared import Pt, RGBColor

    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    for row in rows:
        while len(row) < n_cols:
            row.append("")

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"

            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(30, 41, 59)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EDF5"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            is_total_row = any("total" in c.lower() for c in row_data) and r_idx == len(rows) - 1
            if is_total_row:
                run.bold = True
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4FF"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph("")


@app.get("/api/export/{record_id}/docx")
async def export_docx(record_id: int):
    from .database import GenerationHistory, get_session

    with get_session() as session:
        record = session.get(GenerationHistory, record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Install with: pip install python-docx")

    org_name = get_setting("org_name") or "Organization"
    logo_path_str = get_setting("logo_path")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(30, 41, 59)

    if logo_path_str and Path(logo_path_str).exists():
        try:
            doc.add_picture(logo_path_str, height=Inches(0.8))
        except Exception:
            pass

    title = doc.add_heading(org_name, level=0)
    title.runs[0].font.color.rgb = RGBColor(30, 41, 59)

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Skill: {record.skill_name}  |  Generated: {record.created_at.strftime('%B %d, %Y')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("")

    content = record.output_content or ""
    all_lines = content.split("\n")
    i = 0
    while i < len(all_lines):
        stripped = all_lines[i].strip()

        if stripped.startswith("|") and i + 1 < len(all_lines) and all_lines[i + 1].strip().startswith("|"):
            rows, i = _parse_md_table(all_lines, i)
            _render_docx_table(doc, rows)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        elif stripped == "":
            pass
        else:
            doc.add_paragraph(stripped)
        i += 1

    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Generated by ZettaBrain  |  {record.created_at.strftime('%B %d, %Y at %I:%M %p')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    from fastapi.responses import FileResponse

    safe_name = record.skill_name.replace(" ", "_")
    return FileResponse(
        tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{safe_name}_{record.created_at.strftime('%Y%m%d')}.docx",
    )


@app.get("/api/corpus/summary")
async def corpus_summary():
    """Return a summary of ingested documents for the welcome screen."""
    sources = _get_sources()

    if not sources:
        return {"has_docs": False, "summary": None, "doc_count": 0, "sources": []}

    skills = []
    for skill_dir in [SKILLS_DIR, _BUILTIN_SKILLS_DIR]:
        if skill_dir.exists():
            for f in skill_dir.iterdir():
                if f.suffix == ".md":
                    try:
                        import frontmatter

                        fm = frontmatter.load(str(f))
                        skills.append(
                            {
                                "name": fm.get("name", f.stem),
                                "description": fm.get("description", ""),
                            }
                        )
                    except Exception:
                        pass

    return {
        "has_docs": True,
        "doc_count": len(sources),
        "sources": [Path(s).stem for s in sources[:10]],
        "skills": skills,
    }


# ── Routes: Document Picker (for wizard example) ───────────────────────────
@app.get("/api/documents")
async def list_documents():
    """Return list of ingested documents with metadata for the wizard document picker."""
    if not INGEST_LOG.exists():
        return {"documents": []}
    try:
        data = json.loads(INGEST_LOG.read_text(encoding="utf-8"))
    except Exception:
        return {"documents": []}

    documents = []
    for filepath, file_hash in data.items():
        p = Path(filepath)
        name = p.name
        ext = p.suffix.lstrip(".").lower()
        size_kb = 0
        if p.exists():
            size_kb = round(p.stat().st_size / 1024)
        documents.append({"path": filepath, "name": name, "ext": ext, "size_kb": size_kb})
    documents.sort(key=lambda d: d["name"].lower())
    return {"documents": documents}


@app.get("/api/documents/content")
async def get_document_content(path: str):
    """Read a document's text content for use as a skill example. Supports txt, md, docx, pdf."""
    if not INGEST_LOG.exists():
        raise HTTPException(status_code=404, detail="No ingested documents found")

    try:
        data = json.loads(INGEST_LOG.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(status_code=500, detail="Could not read document index")

    if path not in data:
        raise HTTPException(status_code=404, detail="Document not found in ingested files")

    p = Path(path)
    if not p.exists():
        raise HTTPException(status_code=404, detail="Document file no longer exists on disk")

    ext = p.suffix.lower()
    try:
        if ext in (".txt", ".md"):
            content = p.read_text(encoding="utf-8", errors="replace")
        elif ext == ".docx":
            try:
                from docx import Document
            except ImportError:
                raise HTTPException(status_code=500, detail="python-docx is required to read .docx files")
            doc = Document(str(p))
            content = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        elif ext == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                raise HTTPException(status_code=500, detail="pypdf is required to read .pdf files")
            reader = PdfReader(str(p))
            pages = [page.extract_text() or "" for page in reader.pages]
            content = "\n\n".join(pages)
        else:
            content = p.read_text(encoding="utf-8", errors="replace")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=500, detail="Could not read document content")

    if len(content) > 8000:
        content = content[:8000] + "\n\n[... truncated for preview ...]"

    return {"path": path, "name": p.name, "content": content}


# ── Routes: Clear Vectorstore ────────────────────────────────────────────────
@app.delete("/api/vectorstore")
async def clear_vectorstore():
    _reset_vs_cache()
    try:
        import chromadb

        chromadb.PersistentClient(path=str(CHROMA_PATH)).delete_collection("zettabrain_docs")
    except Exception:
        if CHROMA_PATH.exists():
            shutil.rmtree(str(CHROMA_PATH), ignore_errors=True)
            CHROMA_PATH.mkdir(parents=True, exist_ok=True)

    if INGEST_LOG.exists():
        INGEST_LOG.write_text("{}", encoding="utf-8")

    return {"success": True, "message": "Vector store cleared."}


# ── Helpers: Skills ──────────────────────────────────────────────────────────
def _find_skill_file(skill_name: str):
    from .generation.skill_parser import SkillParser

    for skills_dir in [SKILLS_DIR, _BUILTIN_SKILLS_DIR]:
        if not skills_dir.exists():
            continue
        for f in skills_dir.glob("*.md"):
            try:
                skill = SkillParser.parse_file(f)
                if skill.name == skill_name:
                    return f
            except Exception:
                continue
    return None


def _build_corpus_retriever():
    class LiteCorpusRetriever:
        def get_context_for_generation(self, query, n_results=5, min_relevance=0.3, **kwargs):
            try:
                vs = _get_vs()
                sources = hybrid_retrieve(query, vs, top_k=n_results)
                if not sources:
                    return None, []

                from dataclasses import dataclass

                @dataclass
                class Citation:
                    document_title: str
                    citation_ref: str = ""

                context_parts = ["# CORPUS CONTEXT (from document library)"]
                citations = []
                for i, doc in enumerate(sources, 1):
                    source = Path(doc.metadata.get("source", "unknown")).stem
                    context_parts.append(f"\n## Source [{i}]: {source}")
                    context_parts.append(doc.page_content)
                    citations.append(Citation(document_title=source, citation_ref=f"[{i}]"))

                return "\n".join(context_parts), citations
            except Exception:
                return None, []

    return LiteCorpusRetriever()
